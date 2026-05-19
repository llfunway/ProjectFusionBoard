# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("docs") / "vertical_corr_nT_algorithm_chain.docx"


def set_run_font(run, size=10.5, bold=False, italic=False, color=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def normal_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, size=9.5, color=(40, 40, 40), font="Consolas")
    return p


def add_math_paragraph(doc, formula):
    """Insert an editable Word math object using Office Math XML text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = formula
    mr.append(mt)
    omath.append(mr)
    p._p.append(omath)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_doc_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        s = doc.styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(0, 0, 0)


def build_doc():
    OUT.parent.mkdir(exist_ok=True)

    doc = Document()
    set_doc_style(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("vertical_corr_nT 算法链路说明")
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ProjectFusionBoard 周报整理版")
    set_run_font(r, color=(90, 90, 90))

    normal_para(
        doc,
        "本文整理当前板级 vertical_corr_nT 的实时算法链路。该链路以三轴磁场和 STM32 Mahony 姿态角为输入，"
        "在不依赖完整磁三轴椭球标定 A,bias 和磁轴-IMU 安装矩阵 R 的前提下，通过垂向投影、"
        "RLS 姿态相关残差补偿以及二状态 Kalman/rate 滤波输出修正后的垂向磁场。"
    )

    doc.add_heading("1. 总体链路", level=1)
    normal_para(doc, "当前 vertical_corr_nT 的处理流程如下：")
    add_code_block(
        doc,
        "AD7177 三轴磁场采集\n"
        "  -> 电压到 nT 换算\n"
        "  -> MPU6050 + Mahony 解算 roll/pitch\n"
        "  -> 计算地理垂向单位向量\n"
        "  -> vertical_raw_nT 垂向投影\n"
        "  -> RLS 姿态相关残差补偿\n"
        "  -> 二状态 Kalman / rate filter 收尾\n"
        "  -> vertical_corr_nT"
    )

    doc.add_heading("2. 三轴磁场换算", level=1)
    normal_para(doc, "三个 AD7177 通道采集磁传感器电压，板级先乘以前端比例系数，再按各轴零点和灵敏度换算成 nT。")
    add_math_paragraph(doc, "V_i = G_frontend · ADC_i")
    add_math_paragraph(doc, "B_i = ((V_i - V_0,i) / S_i) · 100000")
    normal_para(doc, "其中 B_i 对应 b1_nT、b2_nT、b3_nT，单位为 nT；S_i 为各轴 V/Oe 灵敏度，100000 为 Oe 到 nT 的换算系数。")

    doc.add_heading("3. 姿态角和垂向单位向量", level=1)
    normal_para(doc, "MPU6050 输出加速度和角速度，STM32 端使用 Mahony 四元数算法得到 roll 和 pitch。当前垂向投影只使用 roll/pitch，不使用 yaw。")
    add_math_paragraph(doc, "φ = roll,    θ = pitch")
    normal_para(doc, "地理垂向在传感器坐标系下的单位向量为：")
    add_math_paragraph(doc, "v = [-sin(θ),  sin(φ) cos(θ),  cos(φ) cos(θ)]^T")
    normal_para(doc, "该向量描述了当前姿态下“地理垂向”落在磁传感器三轴上的方向余弦。")

    doc.add_heading("4. vertical_raw_nT 垂向投影", level=1)
    normal_para(doc, "将三轴磁场向量投影到垂向单位向量上，得到未经残差补偿的垂向磁场 vertical_raw_nT。")
    add_math_paragraph(doc, "B = [B_x, B_y, B_z]^T")
    add_math_paragraph(doc, "vertical_raw_nT = B^T v = B_x v_x + B_y v_y + B_z v_z")
    normal_para(doc, "这一步是整个链路的物理基础。其误差来源包括姿态角误差、磁三轴比例和零点误差、磁轴与 IMU 轴不一致、采样延迟以及动态运动下的加速度污染。")

    doc.add_heading("5. RLS 姿态相关残差补偿", level=1)
    normal_para(doc, "RLS 部分用于补偿与姿态摆动和水平磁场耦合相关的残差。算法先从 vertical_raw_nT 中提取慢变趋势，再将 raw 与趋势之间的差值作为待学习残差。")
    add_math_paragraph(doc, "trend_k = trend_(k-1) + α_trend (vertical_raw,k - trend_(k-1))")
    add_math_paragraph(doc, "target_residual_k = vertical_raw,k - trend_k")
    normal_para(doc, "当前板级版本使用水平磁场分量作为 RLS 特征。水平分量通过从三轴磁场中扣除垂向分量得到：")
    add_math_paragraph(doc, "B_horizontal,k = B_k - vertical_raw,k · v_k")
    add_math_paragraph(doc, "x_k = B_horizontal,k / 50000 = [h_x, h_y, h_z]^T")
    normal_para(doc, "RLS 在线估计残差模型权重：")
    add_math_paragraph(doc, "residual_hat_k = w_k^T x_k")
    add_math_paragraph(doc, "e_k = target_residual_k - w_(k-1)^T x_k")
    add_math_paragraph(doc, "K_k = P_(k-1) x_k / (λ + x_k^T P_(k-1) x_k)")
    add_math_paragraph(doc, "w_k = clamp(w_(k-1) + K_k e_k)")
    add_math_paragraph(doc, "P_k = (P_(k-1) - K_k x_k^T P_(k-1)) / λ")
    normal_para(doc, "完成 RLS 估计后，从原始垂向投影中扣除估计残差：")
    add_math_paragraph(doc, "vertical_rls,k = vertical_raw,k - residual_hat_k")

    doc.add_heading("6. Kalman / rate filter 收尾", level=1)
    normal_para(doc, "RLS 输出仍可能存在短时噪声和毛刺，因此后端使用二状态 Kalman/rate filter 做平滑收尾。状态向量包括垂向磁场值和垂向变化率。")
    add_math_paragraph(doc, "x_k = [vertical_k,  vertical_rate_k]^T")
    add_math_paragraph(doc, "F = [[1, dt], [0, 1]],    H = [1, 0]")
    normal_para(doc, "预测与更新过程为：")
    add_math_paragraph(doc, "x_k^- = F x_(k-1)")
    add_math_paragraph(doc, "P_k^- = F P_(k-1) F^T + Q")
    add_math_paragraph(doc, "innovation_k = vertical_rls,k - H x_k^-")
    add_math_paragraph(doc, "K_k = P_k^- H^T / (H P_k^- H^T + R)")
    add_math_paragraph(doc, "x_k = x_k^- + K_k innovation_k")
    normal_para(doc, "为了减小滤波相位滞后，最终输出加入了很小的 rate lead：")
    add_math_paragraph(doc, "vertical_corr_nT = vertical_k + T_lead · vertical_rate_k")

    doc.add_heading("7. 当前板级参数", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    set_cell_text(headers[0], "参数", bold=True)
    set_cell_text(headers[1], "当前值", bold=True)
    set_cell_text(headers[2], "作用", bold=True)
    rows = [
        ("MAG_FILTER_SAMPLE_DT_S", "0.0200 s", "算法采样周期，对应 50 Hz"),
        ("MAG_FEATURE_SCALE_NT", "50000 nT", "RLS 特征归一化尺度"),
        ("MAG_TREND_CUTOFF_HZ", "0.2000 Hz", "慢趋势低通截止频率"),
        ("MAG_RLS_LAMBDA", "0.9990", "RLS 遗忘因子"),
        ("MAG_RLS_INITIAL_P", "20.0", "RLS 初始协方差"),
        ("MAG_RLS_MAX_WEIGHT_NT", "2500 nT", "RLS 权重限幅"),
        ("MAG_RATE_PROCESS_ACCEL_NT_S2", "900 nT/s^2", "rate Kalman 过程噪声强度"),
        ("MAG_RATE_MEAS_NOISE_NT", "180 nT", "rate Kalman 观测噪声"),
        ("MAG_RATE_LEAD_S", "0.1200 s", "输出 rate lead，降低相位滞后"),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], a)
        set_cell_text(cells[1], b)
        set_cell_text(cells[2], c)

    doc.add_heading("8. 输出含义", level=1)
    add_bullet(doc, "vertical_raw_nT：三轴磁场按照当前 roll/pitch 直接投影到地理垂向后的结果，是未补偿的垂向磁场。")
    add_bullet(doc, "vertical_corr_nT：vertical_raw_nT 经 RLS 姿态相关残差补偿和二状态 Kalman/rate filter 收尾后的结果，是当前板级 deployed 输出。")
    add_bullet(doc, "correction_nT：RLS 根据水平磁场特征估计出的姿态相关残差补偿量。")
    add_bullet(doc, "residual_nT：vertical_raw_nT 与 corrected 输出之间的剩余差值，用于观察补偿量和滤波效果。")

    doc.add_heading("9. 当前方案特点与局限", level=1)
    normal_para(doc, "当前方案的优点是计算量小、状态量少、可在 STM32 上实时部署，并且不依赖完整离线标定数据。它对尖峰毛刺和一部分中高频姿态扰动有明显抑制作用。")
    normal_para(doc, "但该方案仍不能完全替代物理标定。低频大角度摆动下，姿态投影误差、磁轴-IMU 安装角误差、三轴比例和零偏误差会被放大为低频垂向残差，而这类残差在频域上与真实磁异常相似，不能简单通过 drift removal 或低通滤波硬消除。")
    add_bullet(doc, "若不做 A,bias 和 R 标定，低频大角度扰动下很难稳定达到 200 nT。")
    add_bullet(doc, "RLS 只能学习与姿态和水平磁场相关的残差，不能无限制学习所有低频趋势，否则会吞掉真实磁异常。")
    add_bullet(doc, "后续若结合被动万向节或小角度机械稳定结构，可显著降低算法面对的姿态扰动幅度。")

    doc.add_heading("10. 阶段性结论", level=1)
    normal_para(doc, "当前 vertical_corr_nT 是一条轻量化板级实时补偿链路：先用 Mahony 姿态进行垂向投影，再用 RLS 补偿姿态相关残差，最后用二状态 Kalman/rate filter 平滑收尾。该链路已具备板级部署能力，适合用于中小幅扰动下的实时垂向磁场校正。若要进一步向 200 nT 目标逼近，需要结合磁三轴标定、磁轴-IMU 安装关系标定、姿态可信度判断以及被动机械稳定结构共同优化。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("代码对应文件：Core/Src/main.c, Core/Src/mag_residual_filter.c, Core/Inc/mag_residual_filter.h")
    set_run_font(r, size=9, italic=True, color=(100, 100, 100))

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build_doc()
